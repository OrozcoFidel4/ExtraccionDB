import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.tree import DecisionTreeClassifier
from sklearn.metrics import accuracy_score
from docx import Document

doc = Document()
doc.add_heading("Informe del modelo de arbol de desicion", level=1)

df = pd.read_csv(r'C:\Users\fidel\Desktop\TURNO VESPERTINO\IDGS9-1\DB\parcial2\practicasPython\clientes.csv')

#Datos del worrd
doc.add_heading("1. Informe de datos",level=2)
doc.add_paragraph("Primeras filas del dataset:")
doc.add_paragraph(str(df.head()))

x = df[['Edad','Ingreso','Genero']]
y = df['Compra']

#division de los datos en entrenamiento y prueba
X_train, X_test, y_train, y_test = train_test_split(x, y, test_size=0.3, random_state=42)

#modelo
clf = DecisionTreeClassifier()
clf.fit(X_train,y_train)

#prediccion y evaluacion
y_pred = clf.predict(X_test)
precision = accuracy_score(y_test, y_pred)

#informe de precision
doc.add_heading("2. Rendimiento del Modelo: ", level=2)
doc.add_paragraph(f'Precision del Modelo: {precision:.2f}')

#prediccion para un nuevo cliente 
nuevoCliente = pd.DataFrame([[50, 35000, 1]], columns=['Edad', 'Ingreso', 'Genero'])
prediccion = clf.predict(nuevoCliente)

#resultado de la prediccion 
doc.add_heading('3. Prediccion para Nuevo Cliente', level=2)
doc.add_paragraph('Informacion del nuevo cliente:')
doc.add_paragraph('Edad: 50, Ingreso: 35000, Genero: Masculino')
doc.add_paragraph(f"Prediccion: {'Compra' if prediccion[0] == 1 else 'No compra'}")

#documento
doc.save('Informe01.docx')
print("El informe se ha generado y guardado como 'Informe01.docx")
