import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.tree import DecisionTreeClassifier
from sklearn.metrics import accuracy_score
from docx import Document

from sklearn.linear_model import LinearRegression
from sklearn.metrics import mean_squared_error

doc = Document()
doc.add_heading("Informe del modelo de arbol de desicion", level=1)

df = pd.read_csv(r'C:\Users\fidel\Desktop\TURNO VESPERTINO\IDGS9-1\DB\parcial2\practicasPython\datos07.csv')

#Datos del word
doc.add_heading("1. Informe de datos",level=2)
doc.add_paragraph("Primeras filas del dataset:")
doc.add_paragraph(str(df.head()))

x = df[['Horas_Estudio','Asistencia','Participacion','calificacion']]
y = df['Aprobo']

#modelo de regesion lineal
modelo = LinearRegression()

#entrenar modelo
modelo.fit(x,y)

#division de los datos en entrenamiento y prueba
X_train, X_test, y_train, y_test = train_test_split(x, y, test_size=0.3, random_state=42)

#modelo
clf = DecisionTreeClassifier()
clf.fit(X_train,y_train)

#prediccion y evaluacion
y_pred = clf.predict(X_test)
precision = accuracy_score(y_test, y_pred)

#informe de precision
doc.add_heading("2. Rendimiento del Modelo: ", level=2)
doc.add_paragraph(f'Precision del Modelo: {precision:.2f}')

#prediccion para un nuevo estudiantes
nuevoEstudiantes = pd.DataFrame([[4, 78, 3, 65], [6, 85, 4, 75], [3, 55, 2, 50], [7, 90, 5, 88], [5, 80, 3, 70]], columns=['Horas_Estudio','Asistencia','Participacion','calificacion'])
prediccion = clf.predict(nuevoEstudiantes)

#ciclo para separar informacion
cont = 0
horas = []
asistencias = []
participaciones = []
calificaciones =[]


for i in range(5):
    
    for columEdad in range(1):
        horas.append(nuevoEstudiantes.iloc[cont, 0]) #0 hace referencia a la columna de horas
    
    for columAntiguedad in range(1):
        asistencias.append(nuevoEstudiantes.iloc[cont, 1]) #1 hace referencia a la columna de asistencias
    
    for columUsoServicio in range(1):
        participaciones.append(nuevoEstudiantes.iloc[cont, 2]) #2 hace referencia a la columna de participacion

    for columSatisfaccion in range(1):
        calificaciones.append(nuevoEstudiantes.iloc[cont, 3]) #3 hace referencia a la columna de calificacion

    cont = cont + 1 


#resultado de la prediccion 
doc.add_heading('3. Prediccion para Nuevo Cliente', level=2)

#Ciclo para imprimir todos los clientes
cont = 0
noEstudiante = 1
for i in range(5):
    doc.add_heading(f'Informacion del Estudiante No. {noEstudiante}', level=5)
    doc.add_paragraph(f'Horas De Estudio: {horas[cont]}\nAsistencia: {asistencias[cont]}\nParticipaciones: {participaciones[cont]}\nCalificacion: {calificaciones[cont]}')
    doc.add_paragraph(f"Prediccion: {'Aprobado' if prediccion[cont] == 1 else 'Reprobado'} \n")
    noEstudiante = noEstudiante + 1 
    cont = cont + 1

#evaluar el modelo
r2 = modelo.score(x,y)
print(f'Coeficiente de determinacion R²: {r2}\n')

#obtener el mse
Y_pred = modelo.predict(x)
mse = mean_squared_error(y, Y_pred)
print(f'Error cuadratico medio (mse): {mse}\n')

#documento
doc.save('Informe03.docx')
print("El informe se ha generado y guardado como 'Informe03.docx") 
