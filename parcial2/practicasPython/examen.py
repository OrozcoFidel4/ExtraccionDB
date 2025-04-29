import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.tree import DecisionTreeClassifier
from sklearn.metrics import accuracy_score
from docx import Document

from sklearn.linear_model import LinearRegression
from sklearn.metrics import mean_squared_error

import matplotlib.pyplot as plt

doc = Document()
doc.add_heading("Informe del modelo de arbol de desicion", level=1)

df = pd.read_csv(r'C:\Users\fidel\Desktop\TURNO VESPERTINO\IDGS9-1\DB\parcial2\practicasPython\DatosExamen.csv')

#Datos del word
doc.add_heading("1. Informe de datos",level=2)
doc.add_paragraph("Primeras filas del dataset:")
doc.add_paragraph(str(df.head()))

x = df[['frencuenciaCardiaca','TiempoReaccion','NivelHidratacion','CalidadDeSueno', 'Resistencia']]
y = df['Condiciones']

#modelo de regesion lineal
modelo = LinearRegression()

#entrenar modelo
modelo.fit(x,y)

#division de los datos en entrenamiento y prueba
X_train, X_test, y_train, y_test = train_test_split(x, y, test_size=0.3, random_state=42)

#modelo
clf = DecisionTreeClassifier()
clf.fit(X_train,y_train)

#prediccion y evaluacion
y_pred = clf.predict(X_test)
precision = accuracy_score(y_test, y_pred)

#informe de precision
doc.add_heading("2. Rendimiento del Modelo: ", level=2)
doc.add_paragraph(f'Precision del Modelo: {precision:.2f}')

#CreacionGrafica
plt.plot(df.iloc[0],label='Deportista 1',marker='o', linestyle="--", color="red")
plt.plot(df.iloc[1],label='Deportista 1',marker='o', linestyle="--", color="blue")
plt.plot(df.iloc[2],label='Deportista 1',marker='o', linestyle="--", color="green")
plt.title("VALORES DE DEPORTISTAS")
plt.legend()
plt.grid(True)
plt.show()


#prediccion para un nuevo deportistas
nuevosDeportistas = pd.DataFrame([[80, 1, 77, 8, 83], [92, 4, 66, 6, 63]], columns=['frencuenciaCardiaca','TiempoReaccion','NivelHidratacion','CalidadDeSueno', 'Resistencia'])
prediccion = clf.predict(nuevosDeportistas) 

#ciclo para separar informacion
cont = 0
frecuecia = []
tiempoReaccion = []
NivelHidratacion = []
CalidadDeSueno =[]
Resistencia = []


for i in range(2):
    
    for columFrecuencia in range(1):
        frecuecia.append(nuevosDeportistas.iloc[cont, 0]) #0 hace referencia a la columna 
    
    for columTiempoReaccion in range(1):
        tiempoReaccion.append(nuevosDeportistas.iloc[cont, 1]) #1 hace referencia a la columna 
    
    for columNivelHidratacion in range(1):
        NivelHidratacion.append(nuevosDeportistas.iloc[cont, 2]) #2 hace referencia a la columna 

    for columCalidadDeSueno in range(1):
        CalidadDeSueno.append(nuevosDeportistas.iloc[cont, 3]) #3 hace referencia a la columna 

    for columnResistencia in range(1):
        Resistencia.append(nuevosDeportistas.iloc[cont, 4]) #4 hace referencia a la columna 

    cont = cont + 1 


#resultado de la prediccion 
doc.add_heading('3. Prediccion para Nuevos Deportistas', level=2)

#Ciclo para imprimir todos los deportistas
cont = 0
noDeportista = 1
for i in range(2):
    doc.add_heading(f'Informacion del Deportista No. {noDeportista}', level=5)
    doc.add_paragraph(f'Frecuencia Cardiaca Promedio: {frecuecia[cont]}\nTiempo De Reaccion: {tiempoReaccion[cont]}\nNivel De Hidratacion: {NivelHidratacion[cont]}\nCalidad De Sueño: {CalidadDeSueno[cont]}\nResistencia: {Resistencia[cont]}')
    doc.add_paragraph(f"Prediccion: {'Alto Rendimiento' if prediccion[cont] == 1 else 'Bajo Rendimiento'} \n")
    noDeportista = noDeportista + 1 
    cont = cont + 1

#evaluar el modelo
r2 = modelo.score(x,y)
print(f'Coeficiente de determinacion R²: {r2}\n')

#obtener el mse
Y_pred = modelo.predict(x)
mse = mean_squared_error(y, Y_pred)
print(f'Error cuadratico medio (mse): {mse}\n')

#documento
doc.save('InformeExamen.docx')
print("El informe se ha generado y guardado como 'InformeExamen.docx") 
