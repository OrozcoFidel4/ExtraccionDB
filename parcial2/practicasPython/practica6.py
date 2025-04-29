import pandas as pd
from sklearn.model_selection import train_test_split
from sklearn.tree import DecisionTreeClassifier
from sklearn.metrics import accuracy_score
from docx import Document

doc = Document()
doc.add_heading("Informe del modelo de arbol de desicion", level=1)

df = pd.read_csv(r'C:\Users\fidel\Desktop\TURNO VESPERTINO\IDGS9-1\DB\parcial2\practicasPython\datos06.csv')

#Datos del word
doc.add_heading("1. Informe de datos",level=2)
doc.add_paragraph("Primeras filas del dataset:")
doc.add_paragraph(str(df.head()))

x = df[['Edad','Antiguedad','Uso_Servicio','Satisfaccion']]
y = df['Renovacion']

#division de los datos en entrenamiento y prueba
X_train, X_test, y_train, y_test = train_test_split(x, y, test_size=0.3, random_state=42)

#modelo
clf = DecisionTreeClassifier()
clf.fit(X_train,y_train)

#prediccion y evaluacion
y_pred = clf.predict(X_test)
precision = accuracy_score(y_test, y_pred)

#informe de precision
doc.add_heading("2. Rendimiento del Modelo: ", level=2)
doc.add_paragraph(f'Precision del Modelo: {precision:.2f}')

#prediccion para un nuevo clientes 
nuevoClientes = pd.DataFrame([[35, 4, 7, 8], [40, 2, 6, 5], [29, 5, 8, 9], [50, 1, 5, 7], [26, 3, 9, 10]], columns=['Edad','Antiguedad','Uso_Servicio','Satisfaccion'])
prediccion = clf.predict(nuevoClientes)

#ciclo para separar informacion
cont = 0
edades = []
antiguedades = []
usoServicios = []
satisfacciones =[]

for i in range(5):
    
    for columEdad in range(1):
        edad = nuevoClientes.iloc[cont, 0] #0 hace referencia a la columna de Edad
        edades.append(edad)
        
    for columAntiguedad in range(1):
        antiguedad = nuevoClientes.iloc[cont, 1] #1 hace referencia a la columna de Antiguedad
        antiguedades.append(antiguedad)
    
    for columUsoServicio in range(1):
        usoServicio = nuevoClientes.iloc[cont, 2] #2 hace referencia a la columna de uso Servicio
        usoServicios.append(usoServicio)

    for columSatisfaccion in range(1):
        satisfaccion = nuevoClientes.iloc[cont, 3] #3 hace referencia a la columna de Satisfaccion
        satisfacciones.append(satisfaccion)

    cont = cont + 1 
        

#resultado de la prediccion 
doc.add_heading('3. Prediccion para Nuevo Cliente', level=2)

#Ciclo para imprimir todos los clientes
cont = 0
noCliente = 1
for i in range(5):
    doc.add_heading(f'Informacion del cliente No. {noCliente}', level=5)
    doc.add_paragraph(f'Edad: {edades[cont]}\nAntiguedad: {antiguedades[cont]}\nUso Servicio: {usoServicios[cont]}\nSatisfaccion: {satisfacciones[cont]}')
    doc.add_paragraph(f"Prediccion: {'Compra' if prediccion[cont] == 1 else 'No compra'} \n")
    noCliente = noCliente + 1 
    cont = cont + 1

#documento
doc.save('Informe02.docx')
print("El informe se ha generado y guardado como 'Informe02.docx") 
